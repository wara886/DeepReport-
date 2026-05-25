"""DOCX export utilities for generated research reports."""

from __future__ import annotations

from html import escape
import re
from pathlib import Path
from typing import Any, Dict, List
import zipfile


def export_markdown_to_docx(
    markdown: str,
    output_path: str | Path,
    title: str = "FinSight Research Report",
    metadata: Dict[str, Any] | None = None,
) -> Path:
    """Export a readable DOCX from a Markdown report.

    The preferred backend is python-docx when available. A small OOXML fallback
    keeps local packaging smoke tests usable in minimal environments.
    """

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        return _export_with_python_docx(
            markdown=markdown,
            output_path=output_path,
            title=title,
            metadata=metadata or {},
        )
    except ModuleNotFoundError:
        return _export_minimal_docx(
            markdown=markdown,
            output_path=output_path,
            title=title,
            metadata=metadata or {},
        )


def _export_with_python_docx(markdown: str, output_path: Path, title: str, metadata: Dict[str, Any]) -> Path:
    from docx import Document

    document = Document()
    document.add_heading(title, level=0)
    if metadata:
        table = document.add_table(rows=0, cols=2)
        for key, value in metadata.items():
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value)
        document.add_paragraph("")

    in_code = False
    for item in _markdown_blocks(markdown):
        kind = item["kind"]
        text = item["text"]
        if kind == "code_toggle":
            in_code = not in_code
            continue
        if in_code:
            document.add_paragraph(text)
        elif kind == "heading":
            document.add_heading(text, level=int(item["level"]))
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        elif kind == "numbered":
            document.add_paragraph(text, style="List Number")
        elif kind == "image":
            document.add_paragraph(f"图表：{text}")
        else:
            document.add_paragraph(text)

    document.save(output_path)
    return output_path


def _export_minimal_docx(markdown: str, output_path: Path, title: str, metadata: Dict[str, Any]) -> Path:
    paragraphs: List[str] = [_paragraph_xml(title, style="Title")]
    for key, value in metadata.items():
        paragraphs.append(_paragraph_xml(f"{key}: {value}"))
    in_code = False
    for item in _markdown_blocks(markdown):
        kind = item["kind"]
        text = item["text"]
        if kind == "code_toggle":
            in_code = not in_code
            continue
        style = ""
        if not in_code and kind == "heading":
            style = f"Heading{item['level']}"
        elif not in_code and kind == "bullet":
            text = f"- {text}"
        elif not in_code and kind == "numbered":
            text = f"1. {text}"
        elif not in_code and kind == "image":
            text = f"图表：{text}"
        paragraphs.append(_paragraph_xml(text, style=style))

    document_xml = _document_xml("\n".join(paragraphs))
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _content_types_xml())
        zf.writestr("_rels/.rels", _rels_xml())
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/_rels/document.xml.rels", _empty_doc_rels_xml())
        zf.writestr("docProps/core.xml", _core_xml(title))
        zf.writestr("docProps/app.xml", _app_xml())
    return output_path


def _markdown_blocks(markdown: str) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    for raw_line in str(markdown or "").splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            blocks.append({"kind": "code_toggle", "text": ""})
            continue
        if not line.strip():
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            blocks.append(
                {
                    "kind": "heading",
                    "level": min(len(heading.group(1)), 4),
                    "text": _strip_inline_markdown(heading.group(2)),
                }
            )
            continue
        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            blocks.append({"kind": "bullet", "text": _strip_inline_markdown(bullet.group(1))})
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            blocks.append({"kind": "numbered", "text": _strip_inline_markdown(numbered.group(1))})
            continue
        image = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image:
            blocks.append({"kind": "image", "text": f"{image.group(1)} ({image.group(2)})"})
            continue
        blocks.append({"kind": "paragraph", "text": _strip_inline_markdown(line)})
    return blocks


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"[*_`]+", "", text)
    return text.strip()


def _paragraph_xml(text: str, style: str = "") -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{escape(style)}"/></w:pPr>' if style else ""
    return f"<w:p>{style_xml}<w:r><w:t xml:space=\"preserve\">{escape(str(text))}</w:t></w:r></w:p>"


def _document_xml(body_xml: str) -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {body_xml}
    <w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
  </w:body>
</w:document>"""


def _content_types_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>"""


def _rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>"""


def _empty_doc_rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>"""


def _core_xml(title: str) -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/">
  <dc:title>{escape(title)}</dc:title>
</cp:coreProperties>"""


def _app_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties">
  <Application>FinSight</Application>
</Properties>"""
